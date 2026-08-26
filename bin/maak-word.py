#!/usr/bin/env python3
"""Zet een juridische pagina uit src/pages om naar een Word-document (.docx).

De Markdown in src/pages is de bron: die staat op de website én is de tekst die een
vereniging ondertekent. Zo kunnen de gepubliceerde en de getekende versie niet uit
elkaar lopen.

    pip install --user python-docx
    python3 bin/maak-word.py src/pages/verwerkersovereenkomst.md
    python3 bin/maak-word.py src/pages/subverwerkers.md uit/subverwerkers.docx

Invulvelden in de vorm [NAAM VERENIGING] worden geel gearceerd, zodat de invuller ze
niet over het hoofd ziet. Wilt u een PDF: open het document in Word of LibreOffice en
kies "Opslaan als PDF" — dat is de versie die u laat tekenen.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
except ModuleNotFoundError:  # pragma: no cover - alleen bij een kale omgeving
    sys.exit("python-docx ontbreekt. Installeer het met: pip install --user python-docx")

ROOD = RGBColor(0xD8, 0x3A, 0x2E)  # --rood uit global.css, voor de koppen
# **vet**, *cursief*, `code`, [tekst](url) en losse <span>-tags uit de Markdown.
INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))")
VELD = re.compile(r"(\[[A-Z0-9ÀÉÊËÍÏÓÖÚÜ][^\]]*\])")  # [NAAM VERENIGING]


def lees_frontmatter(tekst: str) -> tuple[dict[str, str], str]:
    """Splitst de YAML-frontmatter van de inhoud. Alleen platte sleutel-waardeparen."""
    if not tekst.startswith("---\n"):
        return {}, tekst
    eind = tekst.index("\n---\n", 3)
    kop: dict[str, str] = {}
    for regel in tekst[4:eind].splitlines():
        if ":" in regel:
            sleutel, _, waarde = regel.partition(":")
            kop[sleutel.strip()] = waarde.strip().strip("'\"")
    return kop, tekst[eind + 5 :]


def paginanummer_in_voettekst(document: Document, links: str) -> None:
    """Voettekst met een 'pagina X van Y'-veld; python-docx kent daar geen API voor."""
    voet = document.sections[0].footer.paragraphs[0]
    voet.text = links
    voet.alignment = WD_ALIGN_PARAGRAPH.LEFT
    voet.add_run("\t\t")
    for stuk in ("PAGE", "NUMPAGES"):
        veld = OxmlElement("w:fldSimple")
        veld.set(qn("w:instr"), stuk)
        if stuk == "NUMPAGES":
            voet.add_run(" van ")
        voet._p.append(veld)
    for run in voet.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)


def schrijf_inline(alinea, tekst: str) -> None:
    """Zet de Markdown-opmaak binnen één regel om naar runs."""
    tekst = re.sub(r"</?[a-zA-Z][^>]*>", "", tekst)  # <span> e.d. uit de webversie
    tekst = tekst.replace(r"\*", "*")
    for stuk in filter(None, INLINE.split(tekst)):
        vet = cursief = code = False
        if stuk.startswith("**") and stuk.endswith("**"):
            stuk, vet = stuk[2:-2], True
        elif stuk.startswith("*") and stuk.endswith("*"):
            stuk, cursief = stuk[1:-1], True
        elif stuk.startswith("`") and stuk.endswith("`"):
            stuk, code = stuk[1:-1], True
        elif stuk.startswith("[") and "](" in stuk:  # een link, geen invulveld
            label, _, doel = stuk[1:].partition("](")
            doel = doel.rstrip(")")
            # Een e-mail- of webadres hoort in het papieren document uitgeschreven te
            # staan; een blauwe link die je niet kunt aanklikken helpt niemand.
            zichtbaar = doel.split("?")[0].removeprefix("mailto:")
            stuk = label if zichtbaar in label else f"{label} ({zichtbaar})"
        # Invulvelden opvallend maken, ook binnen vette tekst.
        for deel in filter(None, VELD.split(stuk)):
            run = alinea.add_run(deel)
            run.bold = vet
            run.italic = cursief
            if code:
                run.font.name = "Consolas"
            if VELD.fullmatch(deel):
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def voeg_tabel_toe(document: Document, rijen: list[list[str]]) -> None:
    tabel = document.add_table(rows=0, cols=len(rijen[0]))
    tabel.style = "Table Grid"
    tabel.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, rij in enumerate(rijen):
        cellen = tabel.add_row().cells
        for cel, waarde in zip(cellen, rij):
            cel.paragraphs[0].text = ""
            schrijf_inline(cel.paragraphs[0], waarde)
            for run in cel.paragraphs[0].runs:
                run.font.size = Pt(9)
                if index == 0:
                    run.bold = True


def splits_tabelrij(regel: str) -> list[str]:
    return [cel.strip() for cel in regel.strip().strip("|").split("|")]


def bouw(bron: Path, doel: Path) -> None:
    kop, inhoud = lees_frontmatter(bron.read_text(encoding="utf-8"))
    document = Document()
    stijl = document.styles["Normal"]
    stijl.font.name = "Calibri"
    stijl.font.size = Pt(10.5)
    stijl.paragraph_format.space_after = Pt(8)

    titel = kop.get("titel", bron.stem).split("|")[0].strip()
    versie = kop.get("versie")
    datum = kop.get("datum")
    onderschrift = f"{titel} — versie {versie} ({datum})" if versie else titel
    paginanummer_in_voettekst(document, onderschrift)
    document.core_properties.title = titel
    document.core_properties.author = "LightMedia — Sporten.app"

    regels = inhoud.splitlines()
    nummer = 0
    while nummer < len(regels):
        regel = regels[nummer]
        gestript = regel.strip()

        if not gestript:
            nummer += 1
            continue

        if gestript == "---":  # scheiding vóór de bijlagen
            document.add_page_break()
            nummer += 1
            continue

        if gestript.startswith("#"):
            niveau = len(gestript) - len(gestript.lstrip("#"))
            alinea = document.add_heading("", level=min(niveau, 4))
            schrijf_inline(alinea, gestript[niveau:].strip())
            for run in alinea.runs:
                run.font.color.rgb = ROOD if niveau == 1 else RGBColor(0x1A, 0x1A, 0x1A)
            nummer += 1
            continue

        if gestript.startswith("|"):
            rijen = []
            while nummer < len(regels) and regels[nummer].strip().startswith("|"):
                cellen = splits_tabelrij(regels[nummer])
                if not all(re.fullmatch(r":?-{2,}:?", cel) for cel in cellen if cel):
                    rijen.append(cellen)
                nummer += 1
            if rijen:
                voeg_tabel_toe(document, rijen)
                document.add_paragraph()
            continue

        opsomming = re.match(r"^([-*]|\d+\.)\s+(.*)", gestript)
        if opsomming:
            stijlnaam = "List Bullet" if opsomming.group(1) in "-*" else "List Number"
            tekst = opsomming.group(2)
            # Doorlopende regels van hetzelfde punt (de Markdown breekt op 90 tekens).
            while nummer + 1 < len(regels) and regels[nummer + 1].startswith("   "):
                nummer += 1
                tekst += " " + regels[nummer].strip()
            schrijf_inline(document.add_paragraph(style=stijlnaam), tekst)
            nummer += 1
            continue

        alinea_regels = [gestript]
        while nummer + 1 < len(regels) and regels[nummer + 1].strip() and not re.match(
            r"^\s*([#|]|[-*]\s|\d+\.\s|---$)", regels[nummer + 1]
        ):
            nummer += 1
            alinea_regels.append(regels[nummer].strip())
        schrijf_inline(document.add_paragraph(), " ".join(alinea_regels))
        nummer += 1

    doel.parent.mkdir(parents=True, exist_ok=True)
    document.save(doel)
    print(f"{doel} geschreven ({doel.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        sys.exit(__doc__)
    bronbestand = Path(sys.argv[1])
    doelbestand = (
        Path(sys.argv[2]) if len(sys.argv) > 2 else Path("uit") / f"{bronbestand.stem}.docx"
    )
    bouw(bronbestand, doelbestand)
